import io
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.engagement import Engagement
from app.models.question import Question, ResponseType
from app.models.response import Response
from app.models.risk_assessment import RiskAssessment, RiskRating

# ── Palette ────────────────────────────────────────────────────────────────────
_NAVY = RGBColor(0x1F, 0x38, 0x64)
_BLUE = RGBColor(0x2E, 0x75, 0xB6)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_MUTED = RGBColor(0x71, 0x80, 0x96)
_RISK_COLOR = {
    RiskRating.CRITICAL: RGBColor(0xEE, 0x00, 0x00),
    RiskRating.HIGH: RGBColor(0xC0, 0x00, 0x00),
    RiskRating.MEDIUM: RGBColor(0xFF, 0xC0, 0x00),
    RiskRating.LOW: RGBColor(0x70, 0xAD, 0x47),
}


# ── Low-level helpers ──────────────────────────────────────────────────────────

def _cell_bg(cell, hex6: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex6)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _run(
    para,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size_pt: float = 11,
    color: Optional[RGBColor] = None,
):
    r = para.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, bold=True, size_pt=14, color=_NAVY)


def _h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, bold=True, size_pt=12, color=_BLUE)


def _table_header(row, cols: list[str]) -> None:
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        _run(cell.paragraphs[0], col, bold=True, size_pt=10, color=_WHITE)
        _cell_bg(cell, "1F3864")


# ── Section builders ───────────────────────────────────────────────────────────

def _cover(doc: Document, eng: Engagement) -> None:
    # Main title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(60)
    _run(t, "INFORMATION SECURITY", bold=True, size_pt=26, color=_NAVY)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t2, "DUE DILIGENCE REPORT", bold=True, size_pt=26, color=_NAVY)

    # Application name
    app = doc.add_paragraph()
    app.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app.paragraph_format.space_before = Pt(30)
    _run(app, eng.application_name, bold=True, size_pt=18, color=_NAVY)

    # Operating companies
    oc_names = "  |  ".join(oc.name for oc in eng.operating_companies)
    if oc_names:
        oc_p = doc.add_paragraph()
        oc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(oc_p, oc_names, size_pt=11, color=_MUTED)

    # Doc number
    dn = doc.add_paragraph()
    dn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dn.paragraph_format.space_before = Pt(16)
    _run(dn, eng.doc_number, size_pt=11, color=_BLUE)

    # Date
    date_str = datetime.now(timezone.utc).strftime("%d %B %Y")
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(dp, f"Generated: {date_str}", size_pt=10, color=_MUTED)

    # Footer
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(56)
    _run(fp, "Albatha IT — Information Security Team", size_pt=10, color=_MUTED)


def _doc_control(doc: Document) -> None:
    _h1(doc, "Document Control")
    date_str = datetime.now(timezone.utc).strftime("%d %B %Y")
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    _table_header(tbl.rows[0], ["Version", "Date", "Author", "Description"])
    data = ["v1.0", date_str, "IS Team", "Initial release"]
    for i, d in enumerate(data):
        tbl.rows[1].cells[i].paragraphs[0].clear()
        _run(tbl.rows[1].cells[i].paragraphs[0], d, size_pt=10)


def _exec_summary(doc: Document) -> None:
    _h1(doc, "1.  Executive Summary")
    p = doc.add_paragraph()
    _run(
        p,
        "This section will be populated following AI-assisted risk assessment review.",
        italic=True,
        size_pt=11,
        color=_MUTED,
    )


def _risk_section(doc: Document, ra: Optional[RiskAssessment]) -> None:
    _h1(doc, "2.  Risk Assessment")

    if ra is None:
        p = doc.add_paragraph()
        _run(p, "No risk assessment has been completed for this engagement.", italic=True, size_pt=11, color=_MUTED)
        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = "Table Grid"
        _table_header(tbl.rows[0], ["Description", "Rating", "Assigned To", "Mitigation"])
        return

    # Overall rating
    if ra.overall_rating:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(6)
        _run(rp, "Overall Risk Rating:  ", bold=True, size_pt=11)
        _run(rp, ra.overall_rating.value, bold=True, size_pt=11, color=_RISK_COLOR.get(ra.overall_rating))

    # Summary
    if ra.summary:
        _h2(doc, "Summary")
        sp = doc.add_paragraph()
        _run(sp, ra.summary, size_pt=11)

    # Risk register
    _h2(doc, "Risk Register")
    items = sorted(ra.risk_items, key=lambda i: i.order)

    if not items:
        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = "Table Grid"
        _table_header(tbl.rows[0], ["Description", "Rating", "Assigned To", "Mitigation"])
        return

    tbl = doc.add_table(rows=1 + len(items), cols=4)
    tbl.style = "Table Grid"
    _table_header(tbl.rows[0], ["Description", "Rating", "Assigned To", "Mitigation"])

    for idx, item in enumerate(items):
        row = tbl.rows[idx + 1]
        if idx % 2 == 1:
            for c in row.cells:
                _cell_bg(c, "F2F2F2")

        row.cells[0].paragraphs[0].clear()
        _run(row.cells[0].paragraphs[0], item.description or "", size_pt=10)

        row.cells[1].paragraphs[0].clear()
        _run(
            row.cells[1].paragraphs[0],
            item.rating.value,
            bold=True,
            size_pt=10,
            color=_RISK_COLOR.get(item.rating),
        )

        row.cells[2].paragraphs[0].clear()
        _run(row.cells[2].paragraphs[0], ", ".join(item.assigned_to) if item.assigned_to else "", size_pt=10)

        row.cells[3].paragraphs[0].clear()
        _run(row.cells[3].paragraphs[0], item.mitigation or "", size_pt=10)


def _q_block(doc: Document, question, response, files: list) -> None:
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(10)
    qp.paragraph_format.space_after = Pt(2)
    _run(qp, f"Q{question.question_number}.  {question.question_text}", bold=True, size_pt=11)

    if response is None:
        rp = doc.add_paragraph()
        rp.paragraph_format.left_indent = Cm(0.5)
        _run(rp, "[No response provided]", italic=True, size_pt=10, color=_MUTED)
        return

    if question.response_type == ResponseType.TEXT:
        text = response.response_text
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        if text:
            _run(p, text, size_pt=10)
        else:
            _run(p, "[No response provided]", italic=True, size_pt=10, color=_MUTED)

    elif question.response_type in (ResponseType.SINGLE_CHOICE, ResponseType.MULTI_CHOICE):
        opts = response.selected_options or []
        if opts:
            for opt in opts:
                pp = doc.add_paragraph()
                pp.paragraph_format.left_indent = Cm(0.5)
                _run(pp, f"•  {opt}", size_pt=10)
        else:
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Cm(0.5)
            _run(pp, "[No option selected]", italic=True, size_pt=10, color=_MUTED)

    elif question.response_type == ResponseType.FILE_UPLOAD:
        if not files:
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Cm(0.5)
            _run(pp, "[No files uploaded]", italic=True, size_pt=10, color=_MUTED)
        else:
            for f in files:
                if f.mime_type and f.mime_type.startswith("image/"):
                    try:
                        if Path(f.stored_path).exists():
                            doc.add_picture(f.stored_path, width=Cm(12))
                            cap = doc.add_paragraph()
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _run(cap, f.original_filename, italic=True, size_pt=9, color=_MUTED)
                        else:
                            pp = doc.add_paragraph()
                            pp.paragraph_format.left_indent = Cm(0.5)
                            _run(pp, f"[Image not found: {f.original_filename}]", italic=True, size_pt=10, color=_MUTED)
                    except Exception:
                        pp = doc.add_paragraph()
                        pp.paragraph_format.left_indent = Cm(0.5)
                        _run(pp, f"[Image: {f.original_filename}]", italic=True, size_pt=10, color=_MUTED)
                else:
                    pp = doc.add_paragraph()
                    pp.paragraph_format.left_indent = Cm(0.5)
                    _run(pp, f"[Attachment: {f.original_filename} — see uploaded files]", size_pt=10)


def _questionnaire(
    doc: Document,
    eng: Engagement,
    questions: list,
    responses: dict,
    files_by_q: dict,
) -> None:
    _h1(doc, "3.  Due Diligence Questionnaire")

    standard = [q for q in questions if not q.is_ai_addendum]
    addendum = [q for q in questions if q.is_ai_addendum]

    # Group standard questions by section (preserve order)
    seen: list[str] = []
    by_section: dict[str, list] = {}
    for q in standard:
        if q.section not in by_section:
            seen.append(q.section)
            by_section[q.section] = []
        by_section[q.section].append(q)

    for section in seen:
        _h2(doc, section)
        for q in by_section[section]:
            _q_block(doc, q, responses.get(q.id), files_by_q.get(q.id, []))

    if addendum and eng.is_ai_application:
        doc.add_page_break()
        _h1(doc, "4.  AI Application Addendum")
        for q in addendum:
            _q_block(doc, q, responses.get(q.id), files_by_q.get(q.id, []))


# ── Public API ─────────────────────────────────────────────────────────────────

async def generate_export(engagement_id: uuid.UUID, db: AsyncSession) -> bytes:
    """Build Word export for an engagement. Returns raw .docx bytes."""

    result = await db.execute(
        select(Engagement)
        .where(Engagement.id == engagement_id)
        .options(
            selectinload(Engagement.operating_companies),
            selectinload(Engagement.files),
            selectinload(Engagement.risk_assessment).selectinload(RiskAssessment.risk_items),
            selectinload(Engagement.structured_fields),
        )
    )
    eng = result.scalar_one_or_none()
    if eng is None:
        raise ValueError(f"Engagement {engagement_id} not found")

    q_result = await db.execute(select(Question).order_by(Question.question_number))
    questions = list(q_result.scalars().all())

    r_result = await db.execute(
        select(Response).where(Response.engagement_id == engagement_id)
    )
    responses: dict = {r.question_id: r for r in r_result.scalars().all()}

    files_by_q: dict = {}
    for f in eng.files:
        if f.question_id is not None:
            files_by_q.setdefault(f.question_id, []).append(f)

    # Build document
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    _cover(doc, eng)
    doc.add_page_break()
    _doc_control(doc)
    doc.add_page_break()
    _exec_summary(doc)
    doc.add_page_break()
    _risk_section(doc, eng.risk_assessment)
    doc.add_page_break()
    _questionnaire(doc, eng, questions, responses, files_by_q)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
